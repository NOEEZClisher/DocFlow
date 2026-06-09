from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document as WordDocument

from docflow.converters.docx_converter import DocxConverter
from docflow.converters.registry import default_registry


def save_docx(path: Path, document: Any) -> Path:
    document.save(path)
    return path


def test_docx_converter_converts_paragraphs(tmp_path: Path) -> None:
    word_document = WordDocument()
    word_document.add_paragraph("첫 번째 문단입니다.")
    word_document.add_paragraph("두 번째 문단입니다.")
    path = save_docx(tmp_path / "paragraphs.docx", word_document)

    document = DocxConverter().convert(path)

    assert document.kind == "docx"
    assert document.raw_text == "첫 번째 문단입니다.\n\n두 번째 문단입니다."
    assert document.rendered_html is not None
    assert "<p>첫 번째 문단입니다.</p>" in document.rendered_html


def test_docx_converter_converts_heading_1(tmp_path: Path) -> None:
    word_document = WordDocument()
    word_document.add_heading("제출물 검토", level=1)
    path = save_docx(tmp_path / "heading.docx", word_document)

    document = DocxConverter().convert(path)

    assert document.raw_text == "# 제출물 검토"
    assert document.rendered_html is not None
    assert "<h1>제출물 검토</h1>" in document.rendered_html


def test_docx_converter_converts_tables(tmp_path: Path) -> None:
    word_document = WordDocument()
    table = word_document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "상태"
    table.cell(1, 0).text = "초안"
    table.cell(1, 1).text = "완료"
    path = save_docx(tmp_path / "table.docx", word_document)

    document = DocxConverter().convert(path)

    assert document.raw_text == "\n".join(
        [
            "| 항목 | 상태 |",
            "| --- | --- |",
            "| 초안 | 완료 |",
        ]
    )
    assert document.rendered_html is not None
    assert "<table>" in document.rendered_html


def test_registry_selects_docx_converter_for_docx_files() -> None:
    converter = default_registry.get_converter(Path("submission.docx"))

    assert isinstance(converter, DocxConverter)
