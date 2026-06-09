from __future__ import annotations

from pathlib import Path
from typing import Any

from docflow.core.document_model import Document, DocumentLoadError
from docflow.core.markdown_renderer import render_markdown

from .base import DocumentConverter


class DocxConverter(DocumentConverter):
    """Convert DOCX paragraphs, headings, and tables into Markdown."""

    supported_extensions = frozenset({".docx"})

    def convert(self, path: Path) -> Document:
        try:
            from docx import Document as DocxDocument
        except ModuleNotFoundError as exc:
            raise DocumentLoadError("python-docx 패키지가 설치되어 있지 않습니다.") from exc

        try:
            docx_document = DocxDocument(str(path))
        except Exception as exc:
            raise DocumentLoadError(f"DOCX 파일을 읽을 수 없습니다: {path}") from exc

        raw_text = "\n\n".join(_iter_markdown_blocks(docx_document)).strip()
        return Document(
            path=path,
            kind="docx",
            raw_text=raw_text,
            rendered_html=render_markdown(raw_text),
        )


def _iter_markdown_blocks(document: Any) -> list[str]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blocks: list[str] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            block = _paragraph_to_markdown(Paragraph(child, document))
        elif isinstance(child, CT_Tbl):
            block = _table_to_markdown(Table(child, document))
        else:
            block = ""

        if block:
            blocks.append(block)

    return blocks


def _paragraph_to_markdown(paragraph: Any) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    style_name = paragraph.style.name if paragraph.style is not None else ""
    heading_prefix = {
        "Heading 1": "#",
        "Heading 2": "##",
        "Heading 3": "###",
    }.get(style_name)

    if heading_prefix:
        return f"{heading_prefix} {text}"
    return text


def _table_to_markdown(table: Any) -> str:
    rows = [[_cell_text(cell) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    column_count = max(len(row) for row in rows)
    normalized_rows = [_normalize_row(row, column_count) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * column_count
    body_rows = normalized_rows[1:]

    lines = [
        _format_table_row(header),
        _format_table_row(separator),
    ]
    lines.extend(_format_table_row(row) for row in body_rows)
    return "\n".join(lines)


def _normalize_row(row: list[str], column_count: int) -> list[str]:
    return row + [""] * (column_count - len(row))


def _cell_text(cell: Any) -> str:
    text = " ".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
    return _escape_table_cell(text)


def _escape_table_cell(text: str) -> str:
    return text.replace("\\", "\\\\").replace("|", "\\|").replace("\n", " ")


def _format_table_row(row: list[str]) -> str:
    return "| " + " | ".join(row) + " |"
